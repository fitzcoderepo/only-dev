from __future__ import annotations

import json
import os
import re
from pathlib import Path
from docx import Document
from openai import OpenAI

RESUME_PATH = Path(__file__).parent.parent.parent.parent / "resumes" / "resume.docx"

SECTIONS = ["summary", "skills", "experience", "projects"]

SECTION_LABELS = {
    "summary": "Summary",
    "skills": "Technical Skills",
    "experience": "Professional Experience",
    "projects": "Independent Projects",
}

# Styles that are structural (headings, dates, titles) — never overwrite these
_SKIP_STYLES = {"title", "heading 1", "heading 2", "heading 3", "heading 4"}

# Section heading aliases → section key
_HEADING_ALIASES: dict[str, str] = {
    "summary": "summary",
    "technical skills": "skills",
    "skills": "skills",
    "professional experience": "experience",
    "experience": "experience",
    "independent projects": "projects",
    "projects": "projects",
}

_OWNED_SECTIONS = set(SECTIONS)

def _body_indices(paragraphs) -> dict[str, list[int]]:
    """
    Returns {section: [paragraph indices]} for editable body paragraphs only.
    Skips structural paragraphs (headings, dates, blank separators at section boundaries).
    Stops collecting for a section when the next owned-section heading is hit.
    Education and anything after it is never included.
    """
    result: dict[str, list[int]] = {s: [] for s in _OWNED_SECTIONS}
    current: str | None = None

    for i, para in enumerate(paragraphs):
        style = para.style.name.lower()
        text = para.text.strip()
        text_lower = text.lower()

        # Check if this paragraph is a section heading
        if text_lower in _HEADING_ALIASES:
            current = _HEADING_ALIASES[text_lower]
            continue  # heading itself is never a body paragraph

        # Stop collecting entirely once we hit Education
        if text_lower == "education":
            break

        # Skip all structural/heading-styled paragraphs
        if style in _SKIP_STYLES:
            continue

        # Skip blank paragraphs that sit at the very start of a section
        # (the blank line immediately after a heading)
        if current and not text and not result[current]:
            continue

        if current in _OWNED_SECTIONS:
            result[current].append(i)

    # Trim trailing blank paragraphs from each section
    for section in _OWNED_SECTIONS:
        indices = result[section]
        while indices and not paragraphs[indices[-1]].text.strip():
            indices.pop()

    return result


def get_resume_text() -> str:
    doc = Document(RESUME_PATH)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def get_section_diffs(job: dict) -> dict[str, dict]:
    """Returns {section: {original, suggested}} for each resume section."""
    client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
    resume = get_resume_text()

    section_block = "\n".join(
        f'  "{k}": {{"original": "...", "suggested": "..."}}' for k in SECTIONS
    )

    prompt = f"""You are an expert resume coach. Tailor this resume for the job below.

The resume has exactly these sections in this order:
{chr(10).join(f'- {v}' for v in SECTION_LABELS.values())}

Note: The resume also contains an Education section — ignore it entirely. Do not include it in your output.

Return ONLY a JSON object with this exact shape — no markdown fences, no explanation, no extra keys:
{{
{section_block}
}}

For each section:
- Set "original" to the EXACT text of that section as it appears in the resume, using "-" for bullet points, one per line. For "experience" and "projects", omit role titles, company names, and dates — include only the bullet points themselves.
- Set "suggested" to your tailored version, using the same "-" bullet format, same number of bullets per role or item.
- Do NOT add new bullet points or remove existing ones — only reword within the same structure.
- Do NOT invent experience, skills, or claims the candidate does not already have.
- Only reword, reframe, or reorder existing content to mirror the language and priorities of the job description.
- If a section needs no changes, set "original" and "suggested" to identical text.
- Maintain the same approximate length for each section.
- Only suggest changes that meaningfully improve alignment with the job description — do not reword for the sake of rewording. If a bullet already demonstrates a relevant skill or experience, leave it alone.
- You never fabricate skills or experience.

Job Title: {job['title']}
Company: {job['company']}

Job Description:
{job['description_text']}

Resume:
{resume}"""

    response = client.chat.completions.create(
        model="gpt-5.1",
        messages=[
            {
                "role": "system",
                "content": (
                    "You are an expert resume coach who helps candidates present their genuine "
                    "experience compellingly for specific roles. You never fabricate skills or experience. "
                    "You always respond with valid JSON and nothing else."
                ),
            },
            {"role": "user", "content": prompt},
        ],
        max_completion_tokens=3000,
    )

    raw = response.choices[0].message.content
    raw = re.sub(r"^```json|```$", "", raw.strip(), flags=re.MULTILINE).strip()
    return json.loads(raw)


def build_tailored_docx(approved: dict[str, str], job: dict) -> Path:
    """Write approved section text into a new .docx based on the original resume."""
    from docx import Document as Doc
    from docx.oxml.ns import qn
    from lxml import etree

    src = Doc(RESUME_PATH)
    paras = src.paragraphs
    body_map = _body_indices(paras)

    for section, new_text in approved.items():
        indices = body_map.get(section, [])
        if not indices:
            continue

        # Strip leading "-" bullet markers the model may return, split to lines
        lines = []
        for line in new_text.splitlines():
            line = line.strip().lstrip("-").strip()
            if line:
                lines.append(line)

        for j, idx in enumerate(indices):
            para = paras[idx]
            if j < len(lines):
                # Preserve the formatting of the first run, clear the rest
                first_run = para.runs[0] if para.runs else para.add_run("")
                first_run.text = lines[j]
                for extra_run in para.runs[1:]:
                    extra_run.text = ""
            else:
                # More original paragraphs than new lines — remove the element
                para._element.getparent().remove(para._element)

    # Build output path
    company = re.sub(r"[^\w]", "_", job["company"].lower()).strip("_")
    title = re.sub(r"[^\w]", "_", job["title"].lower()).strip("_")
    title = re.sub(r"_+", "_", title)
    out_dir = RESUME_PATH.parent
    out_path = out_dir / f"{company}_{title}.docx"
    src.save(out_path)
    return out_path